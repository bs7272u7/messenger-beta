# 메신저 구조 안내 Word 문서와 다이어그램을 다시 만들 때만 실행하는 보조 스크립트입니다.
# 서비스 동작에는 연결되지 않으므로 app.py 기능을 수정할 때는 건드리지 않아도 됩니다.
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"C:\Users\user\Desktop\Messenger-Beta\documents")
OUT_DIR.mkdir(exist_ok=True)
DIAGRAM_PATH = OUT_DIR / "cloudchatting_system_flow.png"
DOCX_PATH = OUT_DIR / "cloudchatting_작동구조_안내서.docx"


def font(size, bold=False):
    path = r"C:\Windows\Fonts\malgunbd.ttf" if bold else r"C:\Windows\Fonts\malgun.ttf"
    return ImageFont.truetype(path, size)


def rounded_box(draw, box, fill, outline, title, subtitle=""):
    draw.rounded_rectangle(box, radius=24, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    title_box = draw.textbbox((0, 0), title, font=font(30, True))
    draw.text(((x1 + x2 - (title_box[2] - title_box[0])) / 2, y1 + 24), title, font=font(30, True), fill="#10233f")
    if subtitle:
        lines = subtitle.split("\n")
        y = y1 + 73
        for line in lines:
            line_box = draw.textbbox((0, 0), line, font=font(19))
            draw.text(((x1 + x2 - (line_box[2] - line_box[0])) / 2, y), line, font=font(19), fill="#45566b")
            y += 29


def arrow(draw, start, end, color="#4e8dea"):
    draw.line([start, end], fill=color, width=5)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 10), (x - 18, y + 10)], fill=color)


def make_diagram():
    image = Image.new("RGB", (1600, 920), "#f7faff")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((28, 24, 1572, 896), radius=36, fill="#ffffff", outline="#dce7f6", width=3)
    draw.text((72, 62), "Cloud Chatting 작동 구조", font=font(42, True), fill="#0f3b6e")
    draw.text((74, 120), "사용자 화면에서 발생한 요청이 서버 검증, 데이터 저장, 실시간 전달로 이어지는 흐름", font=font(20), fill="#62748a")

    rounded_box(draw, (70, 265, 330, 430), "#e8f2ff", "#8fc0ff", "사용자", "PC · 모바일 브라우저")
    rounded_box(draw, (430, 220, 760, 475), "#eef6ff", "#7faee8", "웹 화면", "index.html · style.css\nscript.js · 채팅 UI")
    rounded_box(draw, (875, 165, 1225, 355), "#eaf7f2", "#7ab89a", "Flask 서버", "app.py · API · 권한 검증\n세션 · 보안 · 관리자 기능")
    rounded_box(draw, (875, 500, 1225, 690), "#f0edff", "#a794df", "Socket.IO", "실시간 메시지 · 읽음\n온라인 상태 · 테마 동기화")
    rounded_box(draw, (1325, 135, 1530, 300), "#fff4df", "#e7b665", "PostgreSQL", "계정 · 메시지\n친구 · 문의 · 신고")
    rounded_box(draw, (1325, 370, 1530, 535), "#fff0f5", "#df8fb0", "Cloudinary", "사진 · 영상 · 파일")
    rounded_box(draw, (1325, 605, 1530, 770), "#eef7ff", "#79aee2", "외부 서비스", "Resend · Web Push\nGitHub 업데이트")

    arrow(draw, (330, 347), (430, 347))
    arrow(draw, (760, 310), (875, 260))
    arrow(draw, (760, 405), (875, 585))
    arrow(draw, (1225, 260), (1325, 220))
    arrow(draw, (1225, 300), (1325, 450))
    arrow(draw, (1225, 610), (1325, 685))
    draw.text((346, 315), "접속 · 입력", font=font(17), fill="#4e6380")
    draw.text((778, 245), "HTTP API", font=font(17), fill="#4e6380")
    draw.text((758, 470), "실시간 이벤트", font=font(17), fill="#4e6380")
    draw.text((1228, 195), "저장 · 조회", font=font(17), fill="#4e6380")
    draw.text((1233, 408), "업로드", font=font(17), fill="#4e6380")
    draw.text((1235, 655), "메일 · 알림", font=font(17), fill="#4e6380")

    draw.rounded_rectangle((105, 748, 1190, 836), radius=16, fill="#f1f6fd")
    draw.text((145, 771), "핵심 흐름", font=font(21, True), fill="#1c568d")
    draw.text((300, 771), "사용자 행동 → 서버 검증 → DB 저장 → Socket.IO 전달 → 상대방 화면 부분 업데이트", font=font(20), fill="#29445f")
    image.save(DIAGRAM_PATH)


def set_run(run, size=11, color="1A2735", bold=False):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14 if level == 1 else 10, after=6)
    run = p.add_run(text)
    set_run(run, 16 if level == 1 else 13, "2E74B5" if level == 1 else "1F4D78", True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    set_run(p.add_run(text))
    return p


def add_bullet(doc, title, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4)
    set_run(p.add_run(title + " "), bold=True, color="1F4D78")
    set_run(p.add_run(text))
    return p


def add_page_number(paragraph):
    run = paragraph.add_run("페이지 ")
    set_run(run, 9, "67788B")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def make_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("CLOUD CHATTING | 시스템 구조 안내"), 9, "67788B")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=10, after=4)
    set_run(title.add_run("Cloud Chatting 작동 구조 안내서"), 24, "0B3F77", True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=16)
    set_run(subtitle.add_run("웹 메신저의 화면, 서버, 데이터, 실시간 통신 흐름"), 11, "62748A")

    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.65))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, after=10)
    set_run(caption.add_run("그림 1. Cloud Chatting의 전체 요청 및 실시간 전달 구조"), 9, "67788B")

    add_heading(doc, "한눈에 보는 핵심 흐름")
    add_body(doc, "Cloud Chatting은 사용자의 브라우저 화면이 Flask 서버에 요청을 보내고, 서버가 권한과 입력값을 확인한 뒤 PostgreSQL에 저장하는 구조입니다. 저장이 끝난 메시지와 상태 변화는 Socket.IO가 참여자의 화면에 즉시 전달합니다.")
    add_bullet(doc, "일반 요청", "로그인, 친구 요청, 프로필 수정, 문의 작성처럼 요청 후 응답을 받는 기능입니다.")
    add_bullet(doc, "실시간 요청", "새 메시지, 읽음 상태, 온라인 상태, 테마 변경처럼 상대방 화면에도 즉시 반영되어야 하는 기능입니다.")
    add_bullet(doc, "파일 처리", "사진, 영상, 음성, 파일은 Cloudinary에 저장하고 DB에는 파일 주소를 보관합니다.")

    add_heading(doc, "1. 사용자 화면: HTML, CSS, JavaScript")
    add_body(doc, "브라우저는 index.html로 화면 뼈대를 만들고, style.css로 디자인과 반응형 화면, 테마를 적용합니다. script.js는 버튼 클릭, 메시지 전송, 모달 열기, 채팅방 테마 변경, 화면의 부분 업데이트를 담당합니다.")
    add_body(doc, "예를 들어 메시지 전송 버튼을 누르면 JavaScript가 대화방 번호와 메시지 내용을 서버로 전달합니다. 서버가 저장을 완료하면 전체 페이지를 새로고침하지 않고, 변경된 메시지 영역만 화면에 추가합니다.")

    add_heading(doc, "2. Flask 서버: app.py")
    add_body(doc, "app.py는 메신저의 중심입니다. 브라우저가 DB에 직접 접근하지 못하게 하고, 사용자의 요청을 검증한 뒤 필요한 기능만 실행합니다.")
    add_bullet(doc, "계정", "회원가입, 로그인, 이메일 인증, 아이디 및 비밀번호 찾기를 처리합니다.")
    add_bullet(doc, "대화", "1:1 채팅, 그룹 채팅, 메시지, 답장, 공감, 전달, 사진·파일·음성 전송을 처리합니다.")
    add_bullet(doc, "운영", "공지사항, 문의, 리뷰, 신고, 이용 제한, 관리자 페이지 기능을 처리합니다.")
    add_bullet(doc, "보안", "세션, CSRF 방지, 요청 제한, 권한 검증, 관리자 추가 검증을 담당합니다.")

    add_heading(doc, "3. PostgreSQL: 데이터가 유지되는 곳")
    add_body(doc, "PostgreSQL은 서버를 재배포하거나 다시 시작해도 유지되어야 하는 정보를 저장합니다. 비밀번호는 원문이 아닌 해시 형태로 저장하는 것이 원칙입니다.")
    add_bullet(doc, "사용자", "아이디, 이메일, 프로필 사진, 배경사진, 소개글, 공개 범위 정보를 저장합니다.")
    add_bullet(doc, "대화방과 메시지", "채팅방 멤버, 공용 테마, 메시지, 공감, 읽음 위치, 고정·알림 설정을 저장합니다.")
    add_bullet(doc, "운영 데이터", "친구 요청, 차단, 신고, 문의 답변, 공지사항, 리뷰 정보를 저장합니다.")

    add_heading(doc, "4. Socket.IO: 메신저다운 실시간성")
    add_body(doc, "일반 웹 요청만 사용하면 상대방이 새로고침하거나 주기적으로 서버에 물어봐야 새 메시지를 볼 수 있습니다. Socket.IO는 브라우저와 서버의 실시간 연결을 유지해, 메시지와 상태 변화를 바로 전달합니다.")
    add_body(doc, "메시지 전송 → 서버 검증 → DB 저장 → 같은 대화방 참여자에게 실시간 이벤트 전달 → 상대방 화면에 새 메시지 추가의 순서로 동작합니다. 테마 변경, 읽음 상태, 그룹 멤버 변화도 같은 방식으로 동기화됩니다.")

    add_heading(doc, "5. 외부 서비스")
    add_bullet(doc, "Cloudinary", "프로필 사진, 배경사진, 대화 이미지·영상·파일을 보관합니다. DB에는 실제 파일 대신 주소를 저장합니다.")
    add_bullet(doc, "Resend", "회원가입 인증과 계정 찾기·비밀번호 재설정에 필요한 이메일 인증 코드를 발송합니다.")
    add_bullet(doc, "Web Push / VAPID", "사용자가 허용한 브라우저에 새 메시지 알림을 보냅니다. HTTPS와 브라우저 설정이 필요합니다.")
    add_bullet(doc, "GitHub", "커밋 기록을 바탕으로 앱의 업데이트 사항을 사용자 화면에 보여줍니다.")

    add_heading(doc, "메시지 하나가 전달되는 예시")
    steps = [
        "사용자가 메시지를 입력하고 전송 버튼을 누릅니다.",
        "JavaScript가 Flask API에 대화방 번호와 메시지를 전달합니다.",
        "Flask가 로그인, 채팅방 참여, 차단·이용 제한 여부를 검증합니다.",
        "통과한 메시지를 PostgreSQL에 저장합니다.",
        "Socket.IO가 대화방 참여자에게 새 메시지 이벤트를 보냅니다.",
        "상대방 브라우저가 전체 새로고침 없이 새 메시지만 추가합니다.",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4)
        set_run(p.add_run(step))

    add_heading(doc, "정리")
    add_body(doc, "Cloud Chatting은 브라우저 UI, Flask 서버, PostgreSQL, Socket.IO를 중심으로 구성되고 Cloudinary·Resend·Web Push·GitHub가 필요한 기능을 보완하는 구조입니다. 이 구조는 지금의 웹 메신저 기능을 안정적으로 운영하고, 이후 PC 앱과 모바일 앱으로 확장할 때도 서버와 DB를 공통 기반으로 재사용할 수 있습니다.")
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    make_diagram()
    make_document()
    print(DOCX_PATH)
